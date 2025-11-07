#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Create Word document with FULL CONTENT of all articles for a specific date
"""

import json
import requests
from bs4 import BeautifulSoup
import re
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from urllib.parse import unquote
import time

def format_date_arabic(date_str):
    """Format date in Arabic"""
    try:
        date_obj = datetime.strptime(date_str, '%Y-%m-%d')
        arabic_months = [
            'يناير', 'فبراير', 'مارس', 'أبريل', 'مايو', 'يونيو',
            'يوليو', 'أغسطس', 'سبتمبر', 'أكتوبر', 'نوفمبر', 'ديسمبر'
        ]
        return f"{date_obj.day} {arabic_months[date_obj.month - 1]} {date_obj.year}"
    except:
        return date_str

def getTypeLabel(article_type):
    """Convert article type to Arabic label"""
    type_labels = {
        'post': 'مقال',
        'video': 'فيديو',
        'liveblog': 'بث مباشر',
        'episode': 'حلقة',
        'gallery': 'معرض صور'
    }
    return type_labels.get(article_type, article_type)

def decode_url(url):
    """Decode URL-encoded Arabic characters to readable format"""
    try:
        return unquote(url, encoding='utf-8')
    except:
        return url

def fetch_article_content(article_url):
    """Fetch full article content from Al Jazeera website"""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'ar,en-US;q=0.7,en;q=0.3',
            'Accept-Encoding': 'gzip, deflate',
            'Connection': 'keep-alive',
            'Upgrade-Insecure-Requests': '1',
        }
        
        response = requests.get(article_url, headers=headers, timeout=10)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Try to find the main article content
        content_selectors = [
            '.wysiwyg--all-content',  # Al Jazeera main content
            '.article-body',         # Alternative content class
            '.post-content',         # Another alternative
            'article .content',      # Generic article content
            '.entry-content',        # WordPress style
        ]
        
        article_content = None
        for selector in content_selectors:
            content_element = soup.select_one(selector)
            if content_element:
                article_content = content_element
                break
        
        if not article_content:
            # Fallback: try to find any div with substantial text content
            content_divs = soup.find_all('div', string=lambda text: text and len(text.strip()) > 200)
            if content_divs:
                article_content = content_divs[0]
        
        if article_content:
            # Clean up the content
            for unwanted in article_content.find_all(['script', 'style', 'nav', 'header', 'footer', 'aside']):
                unwanted.decompose()
            
            # Replace links with their text content (to avoid line breaks around links)
            # This keeps the text but removes the link structure that causes line breaks
            for link in article_content.find_all('a'):
                # Replace the link with just its text, keeping it inline
                link_text = link.get_text()
                link.replace_with(link_text)
            
            # Extract text - use space separator to keep everything inline
            content_text = article_content.get_text(separator=' ', strip=True)
            
            # Fix: ensure numbers stay with preceding text (no space break before numbers)
            # This prevents numbers from being separated
            content_text = re.sub(r'\s+([٠١٢٣٤٥٦٧٨٩0-9]+)', r' \1', content_text)
            
            # Clean up multiple spaces to single space
            content_text = re.sub(r' +', ' ', content_text)
            
            # Now split into reasonable paragraphs based on sentence endings
            # Split by common Arabic sentence endings (. ! ?) but keep them with the sentence
            # This creates natural paragraph breaks without too many line breaks
            sentences = re.split(r'([.!?]\s+)', content_text)
            
            # Rejoin sentences, grouping them into paragraphs (every 3-4 sentences)
            paragraphs = []
            current_para = []
            for i, part in enumerate(sentences):
                if part.strip():
                    current_para.append(part)
                    # Create paragraph break every 3-4 sentences or at natural breaks
                    if len(current_para) >= 4 and part.strip().endswith(('.', '!', '?')):
                        paragraphs.append(''.join(current_para).strip())
                        current_para = []
            
            # Add remaining sentences
            if current_para:
                paragraphs.append(''.join(current_para).strip())
            
            # Join paragraphs with double newline
            content_text = '\n\n'.join(paragraphs)
            content_text = content_text.strip()
            
            return content_text
        else:
            return "عذراً، لم يتم العثور على محتوى المقال الكامل."
            
    except requests.RequestException as e:
        print(f"Error fetching article content: {e}")
        return "عذراً، حدث خطأ في تحميل محتوى المقال."
    except Exception as e:
        print(f"Error parsing article content: {e}")
        return "عذراً، حدث خطأ في معالجة محتوى المقال."

def create_word_document(articles, date):
    """Create a Word document with FULL CONTENT of articles from a specific date"""
    doc = Document()
    
    # Set document direction to RTL
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Inches(0.12)
    
    # Add title
    title = doc.add_heading(f'أخبار فلسطين - {format_date_arabic(date)}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add summary
    summary = doc.add_paragraph(f'عدد المقالات: {len(articles)} مقال')
    summary.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add separator
    doc.add_paragraph('=' * 50)
    
    # Add articles with full content
    for i, article in enumerate(articles, 1):
        print(f"Processing article {i}/{len(articles)}: {article.get('title', '')[:50]}...")
        
        # Article number and title
        article_heading = doc.add_heading(f'{i}. {article.get("title", "بدون عنوان")}', level=1)
        article_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Article metadata
        metadata = doc.add_paragraph()
        metadata.add_run(f'نوع المحتوى: {getTypeLabel(article.get("type", "post"))}').bold = True
        metadata.add_run(f' | تاريخ النشر: {article.get("date", "غير محدد")}')
        metadata.add_run(f' | المصدر: {article.get("source", "الجزيرة نت")}')
        metadata.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Article excerpt
        if article.get('excerpt'):
            excerpt_heading = doc.add_heading('ملخص المقال:', level=2)
            excerpt_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            excerpt_para = doc.add_paragraph(article['excerpt'])
            excerpt_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Full content
        if article.get('link'):
            content_heading = doc.add_heading('المحتوى الكامل:', level=2)
            content_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            print(f"  Fetching content from: {article.get('link')[:80]}...")
            full_content = fetch_article_content(article['link'])
            
            if full_content and not full_content.startswith('عذراً'):
                # Split content into paragraphs
                paragraphs = full_content.split('\n')
                for para in paragraphs:
                    if para.strip():
                        # Clean up the paragraph text to prevent number displacement
                        cleaned_para = para.strip()
                        # Remove any leading/trailing spaces that might cause issues
                        cleaned_para = re.sub(r'^\s+|\s+$', '', cleaned_para)
                        # Ensure numbers stay with their context (no extra spaces before numbers)
                        cleaned_para = re.sub(r'\s+(\d+)', r' \1', cleaned_para)  # Single space before numbers
                        cleaned_para = re.sub(r'(\d+)\s+', r'\1 ', cleaned_para)  # Single space after numbers
                        
                        content_para = doc.add_paragraph(cleaned_para)
                        content_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        # Set paragraph to keep together to prevent number displacement
                        content_para.paragraph_format.keep_together = True
            else:
                no_content = doc.add_paragraph('لم يتم العثور على المحتوى الكامل')
                no_content.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Article link - DECODED for readability
        link_para = doc.add_paragraph()
        link_para.add_run('رابط المقال الأصلي: ').bold = True
        # Decode the URL to show Arabic characters instead of percent-encoded
        decoded_url = decode_url(article.get('link', 'غير متوفر'))
        link_para.add_run(decoded_url)
        link_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add separator between articles
        if i < len(articles):
            doc.add_paragraph('-' * 50)
            doc.add_paragraph()  # Empty line
    
    # Add footer
    doc.add_paragraph()
    footer = doc.add_paragraph(f'تم إنشاء هذا التقرير في: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc

# Main execution
if __name__ == "__main__":
    date = "2025-11-04"
    
    print(f"Loading articles for date: {date}")
    print("=" * 60)
    
    # Load articles from JSON
    with open('articles_combined.json', 'r', encoding='utf-8') as f:
        all_articles = json.load(f)
    
    # Filter articles by date
    articles_for_date = [article for article in all_articles if article.get('date') == date]
    
    print(f"Found {len(articles_for_date)} articles for {date}")
    print("=" * 60)
    print("\n⚠️  This will take some time as we need to fetch content from each article URL...")
    print("Please be patient...\n")
    
    if not articles_for_date:
        print("No articles found for this date!")
    else:
        # Create Word document with full content
        doc = create_word_document(articles_for_date, date)
        
        # Save to file (with timestamp to avoid conflicts)
        timestamp = datetime.now().strftime("%H%M%S")
        filename = f"palestine_news_full_{date.replace('-', '_')}_{timestamp}.docx"
        doc.save(filename)
        
        print("\n" + "=" * 60)
        print(f"✅ Successfully created Word document with FULL CONTENT!")
        print(f"📁 File saved as: {filename}")
        print(f"📄 Total articles: {len(articles_for_date)}")
        print("🔗 URLs are now displayed in readable format (decoded)")
        print("=" * 60)

