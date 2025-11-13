#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Palestine News Research Web Application
Beautiful web interface for researching the Palestine news dataset
"""

import io
import json
import re
import time
import threading
import uuid
from datetime import datetime, timedelta
from collections import Counter, defaultdict
import os
import requests
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from flask import Flask, render_template, request, jsonify, send_file, Response, stream_with_context, session, redirect, url_for
from pathlib import Path
from urllib.parse import urlparse, unquote, quote
from typing import Optional, Tuple, Callable

from PIL import Image
from bs4 import BeautifulSoup

app = Flask(__name__)
app.secret_key = 'your-secret-key-change-this-in-production'  # For session support

# Translation dictionary
TRANSLATIONS = {
    'ar': {
        'app_name': 'أرشيف أخبار فلسطين',
        'home': 'الرئيسية / البحث',
        'headlines': 'عناوين حسب التاريخ',
        'images': 'معرض الصور',
        'tools': 'أدوات',
        'search_archive': 'ابحث في الأرشيف',
        'search_placeholder': 'اكتب كلمات البحث هنا...',
        'content_type': 'نوع المحتوى',
        'all_types': 'جميع الأنواع',
        'articles': 'مقالات',
        'videos': 'فيديوهات',
        'live_blog': 'بث مباشر',
        'episodes': 'حلقات',
        'search_type': 'نوع البحث',
        'title_and_content': 'العنوان والمحتوى',
        'title_only': 'العنوان فقط',
        'content_only': 'المحتوى فقط',
        'from_date': 'من تاريخ',
        'to_date': 'إلى تاريخ',
        'search': 'بحث',
        'quick_keywords': 'كلمات مفتاحية سريعة',
        'total_articles': 'مقال إخباري',
        'article': 'مقال',
        'video': 'فيديو',
        'liveblog': 'بث مباشر',
        'episode': 'حلقة',
        'gallery': 'معرض صور',
        'search_results': 'نتائج البحث',
        'no_results': 'لم يتم العثور على نتائج',
        'try_different': 'جرب كلمات بحث مختلفة',
        'read_more': 'قراءة المزيد',
        'original_article': 'المقال الأصلي',
        'articles_from': 'مقال إخباري من الجزيرة نت',
        'comprehensive_collection': 'مجموعة شاملة من',
        'all_rights_reserved': 'جميع الحقوق محفوظة',
        'previous': 'السابق',
        'next': 'التالي',
        'searching': 'جاري البحث...',
        'error_occurred': 'حدث خطأ أثناء البحث. يرجى المحاولة مرة أخرى.',
        'please_enter_search': 'يرجى إدخال كلمات البحث',
        'copyright': '&copy; 2025 أرشيف أخبار فلسطين - جميع الحقوق محفوظة',
        'back_to_search': 'العودة للبحث',
        'article_summary': 'ملخص المقال:',
        'full_content': 'المحتوى الكامل:',
        'load_full_content': 'تحميل المحتوى الكامل',
        'click_to_load': 'انقر على "تحميل المحتوى الكامل" لعرض المقال كاملاً',
        'original_link': 'رابط المقال الأصلي:',
        'open_on_aljazeera': 'فتح المقال على الجزيرة نت',
        'article_info': 'معلومات المقال:',
        'source': 'المصدر:',
        'content_type_label': 'نوع المحتوى:',
        'publish_date': 'تاريخ النشر:',
        'article_id': 'معرف المقال:',
        'actions': 'إجراءات:',
        'share_article': 'مشاركة المقال',
        'print_article': 'طباعة المقال',
        'related_articles': 'مقالات ذات صلة',
        'loading': 'جاري التحميل...',
        'cannot_load_related': 'لا يمكن تحميل المقالات ذات الصلة حالياً',
        'no_related_articles': 'لا توجد مقالات ذات صلة',
        'loading_content': 'جاري تحميل المحتوى...',
        'loading_full_content': 'جاري تحميل المحتوى الكامل للمقال...',
        'error_loading_content': 'حدث خطأ في تحميل المحتوى. يرجى المحاولة مرة أخرى.',
        'link_copied': 'تم نسخ رابط المقال إلى الحافظة',
        'choose_date': 'اختر تاريخاً لعرض جميع العناوين من ذلك اليوم',
        'select_date': 'اختر التاريخ',
        'date': 'التاريخ',
        'show_headlines': 'عرض العناوين',
        'export_headlines_only': 'تصدير إلى Word (العناوين فقط)',
        'export_full_content': 'تصدير إلى Word (المحتوى الكامل)',
        'headlines_by_date': 'عناوين الأخبار حسب التاريخ',
        'headlines_title': 'عناوين الأخبار',
        'no_articles_date': 'لا توجد مقالات في هذا التاريخ',
        'try_another_date': 'جرب تاريخاً آخر أو تحقق من صحة التاريخ المحدد',
        'please_select_date': 'يرجى اختيار تاريخ أولاً',
        'exporting': 'جاري التصدير...',
        'tools_dashboard': 'لوحة الأدوات',
        'tools_description': 'إدارة التنزيلات، إنشاء الملفات، والعمل مع الصور بسهولة',
        'create_word_files': 'إنشاء ملفات Word',
        'file_type': 'نوع الملف',
        'word_with_summaries': 'ملف Word مع العناوين والملخصات (كما في صفحة العناوين)',
        'word_headlines_full': 'ملف Word يحتوي على العناوين متبوعة بالمحتوى الكامل (بدون ملخص)',
        'word_with_images': 'ملف Word يشمل الصور مع المحتوى الكامل لكل مقال',
        'creating_file': 'جاري إنشاء الملف...',
        'starting_process': 'بدء العملية...',
        'preparing': 'جاري التجهيز...',
        'cancel': 'إلغاء',
        'download_images': 'تنزيل الصور',
        'date_optional': 'التاريخ (اختياري)',
        'leave_empty_all': 'اتركه فارغاً لتحميل جميع الصور',
        'max_images': 'حد أقصى للصور (اختياري)',
        'force_reload': 'إعادة تحميل الصور حتى لو كانت موجودة (Force)',
        'start_download': 'بدء تنزيل الصور',
        'please_select_date_first': 'يرجى تحديد التاريخ أولاً.',
        'completed': 'اكتمل!',
        'file_created_success': 'تم إنشاء الملف بنجاح! جاري التنزيل...',
        'error_occurred_file': 'حدث خطأ أثناء إنشاء الملف',
        'close': 'إغلاق',
        'connection_error': 'حدث خطأ في الاتصال بالخادم',
        'downloading_images': 'جاري تنزيل الصور ... يرجى الانتظار.',
        'image_gallery': 'معرض الصور الإخباري',
        'gallery_description': 'استعرض المقالات التي تحتوي على صور واكتشف أبرز اللقطات اليومية',
        'results_per_page': 'عدد النتائج في الصفحة',
        'no_articles': 'لا توجد مقالات',
        'article_image': 'صورة المقال',
        'page': 'الصفحة',
        'of': 'من',
        'start_search': 'ابدأ البحث لرؤية النتائج',
        'search_keywords': 'كلمات البحث',
        'search_in_archive': 'البحث في الأرشيف',
        'found_results': 'تم العثور على',
        'result': 'نتيجة',
        'results': 'نتائج',
        'modify_filters': 'قم بتعديل المرشحات',
        'show_results': 'عرض النتائج',
        'reset': 'إعادة التعيين',
        'loading_images': 'جاري تحميل الصور...',
        'no_images_date': 'لا توجد مقالات تحتوي على صور في هذا التاريخ.',
        'view_details': 'عرض التفاصيل',
        'open_image': 'فتح الصورة',
        'showing': 'تم عرض',
        'of': 'من',
        'articles_with_images': 'مقال يحتوي على صور.',
        'error_loading_images': 'حدث خطأ أثناء تحميل الصور. يرجى المحاولة مرة أخرى.',
        'export_success_headlines': 'تم إنشاء الملف مع العناوين والملخصات بنجاح، جاري التنزيل...',
        'export_success': 'تم إنشاء الملف بنجاح، جاري التنزيل...',
        'export_success_images': 'تم إنشاء الملف مع الصور بنجاح، جاري التنزيل...',
        'error_downloading_images': 'حدث خطأ أثناء تنزيل الصور.',
        'requested_images': 'تم طلب',
        'image_with_links': 'صورة تحتوي على روابط.',
        'downloaded_new': 'تم تنزيل',
        'new_image': 'صورة جديدة.',
        'skipped_existing': 'تم تخطي',
        'existing_images': 'صور موجودة مسبقًا.',
        'errors_during': 'أخطاء أثناء التحميل:',
        'folder': 'المجلد:'
    },
    'en': {
        'app_name': 'Palestine News Archive',
        'home': 'Home / Search',
        'headlines': 'Headlines by Date',
        'images': 'Image Gallery',
        'tools': 'Tools',
        'search_archive': 'Search Archive',
        'search_placeholder': 'Enter search keywords here...',
        'content_type': 'Content Type',
        'all_types': 'All Types',
        'articles': 'Articles',
        'videos': 'Videos',
        'live_blog': 'Live Blog',
        'episodes': 'Episodes',
        'search_type': 'Search Type',
        'title_and_content': 'Title and Content',
        'title_only': 'Title Only',
        'content_only': 'Content Only',
        'from_date': 'From Date',
        'to_date': 'To Date',
        'search': 'Search',
        'quick_keywords': 'Quick Keywords',
        'total_articles': 'News Articles',
        'article': 'Article',
        'video': 'Video',
        'liveblog': 'Live Blog',
        'episode': 'Episode',
        'gallery': 'Gallery',
        'search_results': 'Search Results',
        'no_results': 'No results found',
        'try_different': 'Try different search keywords',
        'read_more': 'Read More',
        'original_article': 'Original Article',
        'articles_from': 'news articles from Al Jazeera',
        'comprehensive_collection': 'Comprehensive collection of',
        'all_rights_reserved': 'All rights reserved',
        'previous': 'Previous',
        'next': 'Next',
        'searching': 'Searching...',
        'error_occurred': 'An error occurred while searching. Please try again.',
        'please_enter_search': 'Please enter search keywords',
        'copyright': '&copy; 2025 Palestine News Archive - All rights reserved',
        'back_to_search': 'Back to Search',
        'article_summary': 'Article Summary:',
        'full_content': 'Full Content:',
        'load_full_content': 'Load Full Content',
        'click_to_load': 'Click "Load Full Content" to view the complete article',
        'original_link': 'Original Article Link:',
        'open_on_aljazeera': 'Open Article on Al Jazeera',
        'article_info': 'Article Information:',
        'source': 'Source:',
        'content_type_label': 'Content Type:',
        'publish_date': 'Publish Date:',
        'article_id': 'Article ID:',
        'actions': 'Actions:',
        'share_article': 'Share Article',
        'print_article': 'Print Article',
        'related_articles': 'Related Articles',
        'loading': 'Loading...',
        'cannot_load_related': 'Cannot load related articles at this time',
        'no_related_articles': 'No related articles found',
        'loading_content': 'Loading content...',
        'loading_full_content': 'Loading full article content...',
        'error_loading_content': 'An error occurred while loading content. Please try again.',
        'link_copied': 'Article link copied to clipboard',
        'choose_date': 'Select a date to view all headlines from that day',
        'select_date': 'Select Date',
        'date': 'Date',
        'show_headlines': 'Show Headlines',
        'export_headlines_only': 'Export to Word (Headlines Only)',
        'export_full_content': 'Export to Word (Full Content)',
        'headlines_by_date': 'Headlines by Date',
        'headlines_title': 'Headlines',
        'no_articles_date': 'No articles found for this date',
        'try_another_date': 'Try another date or verify the selected date',
        'please_select_date': 'Please select a date first',
        'exporting': 'Exporting...',
        'tools_dashboard': 'Tools Dashboard',
        'tools_description': 'Manage downloads, create files, and work with images easily',
        'create_word_files': 'Create Word Files',
        'file_type': 'File Type',
        'word_with_summaries': 'Word file with headlines and summaries (as in headlines page)',
        'word_headlines_full': 'Word file containing headlines followed by full content (no summary)',
        'word_with_images': 'Word file including images with full content for each article',
        'creating_file': 'Creating file...',
        'starting_process': 'Starting process...',
        'preparing': 'Preparing...',
        'cancel': 'Cancel',
        'download_images': 'Download Images',
        'date_optional': 'Date (Optional)',
        'leave_empty_all': 'Leave empty to download all images',
        'max_images': 'Maximum Images (Optional)',
        'force_reload': 'Reload images even if they exist (Force)',
        'start_download': 'Start Download',
        'please_select_date_first': 'Please select a date first.',
        'completed': 'Completed!',
        'file_created_success': 'File created successfully! Downloading...',
        'error_occurred_file': 'An error occurred while creating the file',
        'close': 'Close',
        'connection_error': 'A connection error occurred with the server',
        'downloading_images': 'Downloading images... Please wait.',
        'image_gallery': 'News Image Gallery',
        'gallery_description': 'Browse articles with images and discover the most important daily shots',
        'results_per_page': 'Results per page',
        'no_articles': 'No articles found',
        'article_image': 'Article Image',
        'page': 'Page',
        'of': 'of',
        'start_search': 'Start searching to see results',
        'search_keywords': 'Search Keywords',
        'search_in_archive': 'Search in Archive',
        'found_results': 'Found',
        'result': 'result',
        'results': 'results',
        'modify_filters': 'or modify the filters',
        'show_results': 'Show Results',
        'reset': 'Reset',
        'loading_images': 'Loading images...',
        'no_images_date': 'No articles with images found for this date.',
        'view_details': 'View Details',
        'open_image': 'Open Image',
        'showing': 'Showing',
        'of': 'of',
        'articles_with_images': 'articles with images.',
        'error_loading_images': 'An error occurred while loading images. Please try again.',
        'export_success_headlines': 'File with headlines and summaries created successfully, downloading...',
        'export_success': 'File created successfully, downloading...',
        'export_success_images': 'File with images created successfully, downloading...',
        'error_downloading_images': 'An error occurred while downloading images.',
        'requested_images': 'Requested',
        'image_with_links': 'images with links.',
        'downloaded_new': 'Downloaded',
        'new_image': 'new images.',
        'skipped_existing': 'Skipped',
        'existing_images': 'existing images.',
        'errors_during': 'Errors during download:',
        'folder': 'Folder:'
    }
}

# Get current language from session or default to Arabic
def get_language():
    return session.get('language', 'ar')

# Add template function for translations
@app.template_global()
def t(key):
    """Get translation for a key"""
    lang = get_language()
    return TRANSLATIONS.get(lang, TRANSLATIONS['ar']).get(key, key)

# Add template function for type labels
@app.template_global()
def getTypeLabel(article_type):
    """Convert article type to label based on current language"""
    lang = get_language()
    type_key = {
        'post': 'article',
        'video': 'video',
        'liveblog': 'liveblog',
        'episode': 'episode',
        'gallery': 'gallery'
    }.get(article_type, article_type)
    return t(type_key)

# Add custom filter for number formatting
@app.template_filter('number_format')
def number_format(value):
    """Format number with commas"""
    return f"{value:,}"

IMAGES_CACHE_DIR = Path("images_cache")
TEMP_DIR = Path("temp")
TEMP_DIR.mkdir(exist_ok=True)

# Progress tracking for document creation
progress_tracker = {}
progress_lock = threading.Lock()

class NewsAnalyzer:
    def __init__(self, json_file_path):
        """Initialize the analyzer with the combined articles dataset"""
        with open(json_file_path, 'r', encoding='utf-8') as f:
            self.articles = json.load(f)
        print(f"Loaded {len(self.articles)} articles for web analysis")
        self.content_cache = {}  # Cache for fetched article content
    
    def search_articles(self, query, search_type='all', content_type='all', date_from=None, date_to=None):
        """Search articles with multiple filters"""
        results = []
        query_lower = query.lower() if query else ""
        
        for article in self.articles:
            # Text search
            if query:
                title = article.get('title', '').lower()
                excerpt = article.get('excerpt', '').lower()
                
                if search_type == 'title' and query_lower not in title:
                    continue
                elif search_type == 'excerpt' and query_lower not in excerpt:
                    continue
                elif search_type == 'all' and query_lower not in title and query_lower not in excerpt:
                    continue
            
            # Content type filter
            if content_type != 'all' and article.get('type') != content_type:
                continue
            
            # Date range filter
            if date_from and article.get('date', '') < date_from:
                continue
            if date_to and article.get('date', '') > date_to:
                continue
            
            results.append(article)
        
        return results
    
    def get_statistics(self):
        """Get comprehensive statistics"""
        dates = [article.get('date') for article in self.articles if article.get('date')]
        types = [article.get('type') for article in self.articles]
        
        return {
            'total_articles': len(self.articles),
            'date_range': {
                'start': min(dates) if dates else None,
                'end': max(dates) if dates else None
            },
            'article_types': dict(Counter(types)),
            'articles_with_images': len([a for a in self.articles if a.get('image_url')])
        }
    
    def get_timeline_data(self):
        """Get timeline data for charts"""
        monthly_counts = defaultdict(int)
        for article in self.articles:
            date = article.get('date', '')
            if date:
                month = date[:7]  # YYYY-MM
                monthly_counts[month] += 1
        
        return sorted(monthly_counts.items())
    
    def get_keyword_analysis(self, keywords):
        """Analyze keyword frequency"""
        keyword_counts = {}
        for keyword in keywords:
            count = sum(1 for article in self.articles 
                       if keyword in article.get('title', '') or keyword in article.get('excerpt', ''))
            keyword_counts[keyword] = count
        return keyword_counts

    def get_articles_with_images(self, date=None):
        """Return articles that include images, optionally filtered by date"""
        articles_with_images = [
            article for article in self.articles
            if article.get('image_url')
        ]

        if date:
            articles_with_images = [
                article for article in articles_with_images
                if article.get('date') == date
            ]

        # Sort by date (newest first), then by id descending as tie-breaker
        articles_with_images.sort(
            key=lambda article: (
                article.get('date') or '0000-00-00',
                article.get('id') or 0
            ),
            reverse=True
        )

        return articles_with_images
    
    def fetch_article_content(self, article_url):
        """Fetch full article content from Al Jazeera website"""
        # Check cache first
        if article_url in self.content_cache:
            return self.content_cache[article_url]
        
        try:
            # Add headers to mimic a real browser
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
                'Accept-Language': 'ar,en-US;q=0.7,en;q=0.3',
                'Accept-Encoding': 'gzip, deflate',
                'Connection': 'keep-alive',
                'Upgrade-Insecure-Requests': '1',
            }
            
            response = requests.get(article_url, headers=headers, timeout=10)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Try to find the main article content
            content_selectors = [
                '.wysiwyg--all-content',  # Al Jazeera main content
                '.article-body',         # Alternative content class
                '.post-content',         # Another alternative
                'article .content',      # Generic article content
                '.entry-content',        # WordPress style
            ]
            
            article_content = None
            for selector in content_selectors:
                content_element = soup.select_one(selector)
                if content_element:
                    article_content = content_element
                    break
            
            if not article_content:
                # Fallback: try to find any div with substantial text content
                content_divs = soup.find_all('div', string=lambda text: text and len(text.strip()) > 200)
                if content_divs:
                    article_content = content_divs[0]
            
            if article_content:
                # Clean up the content
                # Remove unwanted elements
                for unwanted in article_content.find_all(['script', 'style', 'nav', 'header', 'footer', 'aside']):
                    unwanted.decompose()
                
                # Extract text and format it
                content_text = article_content.get_text(separator='\n', strip=True)
                
                # Clean up extra whitespace
                content_text = re.sub(r'\n\s*\n', '\n\n', content_text)
                content_text = content_text.strip()
                
                # Cache the result
                self.content_cache[article_url] = content_text
                return content_text
            else:
                return "عذراً، لم يتم العثور على محتوى المقال الكامل."
                
        except requests.RequestException as e:
            print(f"Error fetching article content: {e}")
            return "عذراً، حدث خطأ في تحميل محتوى المقال."
        except Exception as e:
            print(f"Error parsing article content: {e}")
            return "عذراً، حدث خطأ في معالجة محتوى المقال."
    
    def create_word_document(self, articles, date, include_content=False, progress_callback: Optional[Callable] = None):
        """Create a Word document with articles from a specific date"""
        doc = Document()
        
        # Set document direction to RTL
        doc.styles['Normal'].font.name = 'Arial'
        doc.styles['Normal'].font.size = Inches(0.12)
        
        total_steps = len(articles) * (2 if include_content else 1) + 3  # per article + init + summary + save
        current_step = 0
        
        def update_progress(step_increment, message, status="processing"):
            nonlocal current_step
            current_step += step_increment
            percentage = int((current_step / total_steps) * 100) if total_steps > 0 else 0
            if progress_callback:
                progress_callback(percentage, message, status)
        
        update_progress(1, f"بدء إنشاء المستند ({len(articles)} مقال)", "processing")
        
        # Add title
        title = doc.add_heading(f'أخبار فلسطين - {self.format_date_arabic(date)}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add summary
        summary = doc.add_paragraph(f'عدد المقالات: {len(articles)} مقال')
        summary.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add separator
        doc.add_paragraph('=' * 50)
        
        update_progress(1, "إضافة العناوين والملخصات...", "processing")
        
        # Add articles
        for i, article in enumerate(articles, 1):
            update_progress(0, f"معالجة المقال {i} من {len(articles)}: {article.get('title', 'بدون عنوان')[:50]}...", "processing")
            
            # Article number and title
            article_heading = doc.add_heading(f'{i}. {article.get("title", "بدون عنوان")}', level=1)
            article_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Article metadata
            metadata = doc.add_paragraph()
            metadata.add_run(f'نوع المحتوى: {getTypeLabel(article.get("type", "post"))}').bold = True
            metadata.add_run(f' | تاريخ النشر: {article.get("date", "غير محدد")}')
            metadata.add_run(f' | المصدر: {article.get("source", "الجزيرة نت")}')
            metadata.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Article excerpt
            if article.get('excerpt'):
                excerpt_heading = doc.add_heading('ملخص المقال:', level=2)
                excerpt_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                excerpt_para = doc.add_paragraph(article['excerpt'])
                excerpt_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Full content if requested
            if include_content and article.get('link'):
                update_progress(0, f"جاري جلب محتوى المقال {i}...", "processing")
                content_heading = doc.add_heading('المحتوى الكامل:', level=2)
                content_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                try:
                    full_content = self.fetch_article_content(article['link'])
                    if full_content and not full_content.startswith('عذراً'):
                        # Split content into paragraphs
                        paragraphs = full_content.split('\n')
                        for para in paragraphs:
                            if para.strip():
                                content_para = doc.add_paragraph(para.strip())
                                content_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        no_content = doc.add_paragraph('لم يتم العثور على المحتوى الكامل')
                        no_content.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                except Exception as e:
                    error_para = doc.add_paragraph(f'خطأ في تحميل المحتوى: {str(e)}')
                    error_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                update_progress(1, f"تمت إضافة محتوى المقال {i}", "processing")
            else:
                update_progress(1, f"اكتمل المقال {i}", "processing")
            
            # Article link (only show if not including full content)
            if not include_content:
                link_para = doc.add_paragraph()
                link_para.add_run('رابط المقال الأصلي: ').bold = True
                link_para.add_run(article.get('link', 'غير متوفر'))
                link_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Add separator between articles
            if i < len(articles):
                doc.add_paragraph('-' * 50)
                doc.add_paragraph()  # Empty line
        
        update_progress(1, "جاري حفظ الملف...", "processing")
        
        # Add footer
        doc.add_paragraph()
        footer = doc.add_paragraph(f'تم إنشاء هذا التقرير في: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        update_progress(1, "اكتمل إنشاء الملف بنجاح!", "completed")
        
        return doc
    
    def format_date_arabic(self, date_str):
        """Format date in Arabic"""
        try:
            date_obj = datetime.strptime(date_str, '%Y-%m-%d')
            arabic_months = [
                'يناير', 'فبراير', 'مارس', 'أبريل', 'مايو', 'يونيو',
                'يوليو', 'أغسطس', 'سبتمبر', 'أكتوبر', 'نوفمبر', 'ديسمبر'
            ]
            return f"{date_obj.day} {arabic_months[date_obj.month - 1]} {date_obj.year}"
        except:
            return date_str

# Initialize analyzer
analyzer = NewsAnalyzer('articles_combined.json')

@app.route('/set_language/<lang>')
def set_language(lang):
    """Set language preference"""
    if lang in ['ar', 'en']:
        session['language'] = lang
    return redirect(request.referrer or url_for('index'))

@app.route('/')
def index():
    """Main page - Direct search interface"""
    stats = analyzer.get_statistics()
    lang = get_language()
    return render_template('index.html', stats=stats, lang=lang)

@app.route('/search')
def search():
    """Search page - redirect to home (search is now on home page)"""
    lang = get_language()
    return render_template('index.html', stats=analyzer.get_statistics(), lang=lang)

@app.route('/headlines')
def headlines():
    """Headlines by date page"""
    lang = get_language()
    return render_template('headlines.html', lang=lang)


@app.route('/images')
def images():
    """Images gallery page"""
    lang = get_language()
    return render_template('images.html', lang=lang)


@app.route('/tools')
def tools():
    """Tools page"""
    lang = get_language()
    return render_template('tools.html', lang=lang)

@app.route('/api/search')
def api_search():
    """API endpoint for searching articles"""
    query = request.args.get('q', '')
    search_type = request.args.get('type', 'all')
    content_type = request.args.get('content_type', 'all')
    date_from = request.args.get('date_from', '')
    date_to = request.args.get('date_to', '')
    page = int(request.args.get('page', 1))
    per_page = int(request.args.get('per_page', 20))
    
    results = analyzer.search_articles(query, search_type, content_type, date_from, date_to)
    
    # Pagination
    total = len(results)
    start = (page - 1) * per_page
    end = start + per_page
    paginated_results = results[start:end]
    
    return jsonify({
        'articles': paginated_results,
        'total': total,
        'page': page,
        'per_page': per_page,
        'total_pages': (total + per_page - 1) // per_page
    })

@app.route('/api/statistics')
def api_statistics():
    """API endpoint for statistics"""
    return jsonify(analyzer.get_statistics())

@app.route('/api/timeline')
def api_timeline():
    """API endpoint for timeline data"""
    return jsonify(analyzer.get_timeline_data())

@app.route('/api/keywords')
def api_keywords():
    """API endpoint for keyword analysis"""
    keywords = request.args.get('keywords', '').split(',')
    keywords = [k.strip() for k in keywords if k.strip()]
    if keywords:
        return jsonify(analyzer.get_keyword_analysis(keywords))
    return jsonify({})

@app.route('/api/headlines')
def api_headlines():
    """API endpoint for headlines by date"""
    date = request.args.get('date', '')
    page = int(request.args.get('page', 1))
    per_page = int(request.args.get('per_page', 50))
    
    # Filter articles by date
    filtered_articles = []
    for article in analyzer.articles:
        if article.get('date') == date:
            filtered_articles.append({
                'id': article.get('id'),
                'title': article.get('title'),
                'link': article.get('link'),
                'type': article.get('type'),
                'date': article.get('date')
            })
    
    # Pagination
    total = len(filtered_articles)
    start = (page - 1) * per_page
    end = start + per_page
    paginated_results = filtered_articles[start:end]
    
    return jsonify({
        'articles': paginated_results,
        'total': total,
        'page': page,
        'per_page': per_page,
        'total_pages': (total + per_page - 1) // per_page
    })


@app.route('/api/images')
def api_images():
    """API endpoint for articles that include images"""
    date = request.args.get('date', '').strip()
    page = int(request.args.get('page', 1))
    per_page = int(request.args.get('per_page', 24))

    # Limit per_page to avoid overly large responses
    per_page = min(max(per_page, 1), 60)

    articles_with_images = analyzer.get_articles_with_images(date or None)

    total = len(articles_with_images)
    start = (page - 1) * per_page
    end = start + per_page
    paginated_articles = articles_with_images[start:end]

    response_articles = [
        {
            'id': article.get('id'),
            'title': article.get('title'),
            'date': article.get('date'),
            'type': article.get('type'),
            'image_url': article.get('image_url'),
            'link': article.get('link'),
            'excerpt': article.get('excerpt'),
        }
        for article in paginated_articles
    ]

    return jsonify({
        'articles': response_articles,
        'total': total,
        'page': page,
        'per_page': per_page,
        'total_pages': (total + per_page - 1) // per_page
    })


@app.route('/api/tools/download-images', methods=['POST'])
def api_download_images():
    data = request.get_json() or {}
    date = data.get('date')
    limit = data.get('limit')
    force = data.get('force', False)

    try:
        limit = int(limit) if limit is not None else None
    except (ValueError, TypeError):
        limit = None

    result = download_images(date=date or None, limit=limit, force=bool(force))

    return jsonify({
        'success': True,
        'requested': result['requested'],
        'downloaded': result['downloaded'],
        'skipped': result['skipped'],
        'errors': result['errors'],
        'output_dir': str(IMAGES_CACHE_DIR.resolve())
    })


def build_image_path(article, output_dir: Path) -> Path:
    image_url = article.get('image_url', '')
    parsed = urlparse(image_url)
    basename = Path(unquote(parsed.path)).name or f"article_{article.get('id', 'unknown')}"
    stem, ext = os.path.splitext(basename)
    if not ext:
        ext = '.jpg'

    article_date = (article.get('date') or 'unknown_date').replace('/', '-').strip()
    target_dir = output_dir / article_date
    target_dir.mkdir(parents=True, exist_ok=True)

    filename = f"{article.get('id', 'article')}_{stem}{ext}"
    return target_dir / filename


def ensure_image_file(article) -> Optional[Path]:
    """Ensure the article image exists locally and return its path."""
    image_url = article.get('image_url')
    if not image_url:
        return None

    destination = build_image_path(article, IMAGES_CACHE_DIR)

    if destination.exists():
        return destination

    try:
        response = requests.get(image_url, timeout=15)
        response.raise_for_status()
        destination.parent.mkdir(parents=True, exist_ok=True)
        destination.write_bytes(response.content)
        return destination
    except requests.RequestException as exc:
        print(f"Failed to download image {image_url}: {exc}")
        return None


def prepare_image_stream(image_path: Path, max_width: int = 1200, max_height: int = 1200,
                         quality: int = 80) -> Optional[io.BytesIO]:
    """Resize/compress image and return BytesIO stream ready for docx embedding."""
    try:
        with Image.open(image_path) as img:
            img = img.convert('RGB')
            img.thumbnail((max_width, max_height), Image.LANCZOS)
            output = io.BytesIO()
            img.save(output, format='JPEG', quality=quality, optimize=True)
            output.seek(0)
            return output
    except Exception as exc:
        print(f"Failed to process image {image_path}: {exc}")
        return None


def fetch_article_content_formatted(article) -> str:
    """Fetch and format full article content; fallback to excerpt if download fails."""
    link = article.get('link')
    if not link:
        return ''

    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
        "Accept-Language": "ar,en-US;q=0.7,en;q=0.3",
        "Connection": "keep-alive",
    }

    try:
        response = requests.get(link, headers=headers, timeout=15)
        response.raise_for_status()

        soup = BeautifulSoup(response.content, 'html.parser')

        selectors = [
            '.wysiwyg--all-content',
            '.article-body',
            '.post-content',
            'article .content',
            '.entry-content',
        ]

        article_content = None
        for selector in selectors:
            node = soup.select_one(selector)
            if node:
                article_content = node
                break

        if not article_content:
            fallback = soup.find_all('div', string=lambda text: text and len(text.strip()) > 200)
            if fallback:
                article_content = fallback[0]

        if not article_content:
            return ''

        for unwanted in article_content.find_all(['script', 'style', 'nav', 'header', 'footer', 'aside']):
            unwanted.decompose()

        for link_node in article_content.find_all('a'):
            link_node.replace_with(link_node.get_text())

        text = article_content.get_text(separator=' \n ', strip=True)
        text = re.sub(r'\s+', ' ', text)

        paragraph_candidates = re.split(r'([.!؟]\s+)', text)
        paragraphs, buffer = [], []
        for part in paragraph_candidates:
            if part.strip():
                buffer.append(part)
                if part.strip().endswith(tuple('.!؟')) and len(buffer) >= 2:
                    paragraphs.append(''.join(buffer).strip())
                    buffer = []
        if buffer:
            paragraphs.append(''.join(buffer).strip())

        if not paragraphs:
            paragraphs = [text]

        return '\n\n'.join(paragraphs)
    except requests.RequestException as exc:
        print(f"Network error while fetching article content {link}: {exc}")
        return ''
    except Exception as exc:
        print(f"Parsing error while fetching article content {link}: {exc}")
        return ''


def download_images(date=None, limit=None, force=False):
    articles = analyzer.get_articles_with_images(date)

    downloaded = 0
    skipped = 0
    errors = 0

    for article in articles:
        if limit and downloaded >= limit:
            break

        destination = build_image_path(article, IMAGES_CACHE_DIR)
        image_url = article.get('image_url')

        if destination.exists() and not force:
            skipped += 1
            continue

        try:
            response = requests.get(image_url, timeout=15)
            response.raise_for_status()
            destination.write_bytes(response.content)
            downloaded += 1
        except requests.RequestException as exc:
            errors += 1
            print(f"Error downloading image {image_url}: {exc}")

    return {
        'requested': len(articles),
        'downloaded': downloaded,
        'skipped': skipped,
        'errors': errors
    }


def create_headline_only_document(date, progress_callback: Optional[Callable] = None):
    """Create document with headlines and full content (same format as with images, but without images)"""
    articles = [
        article for article in analyzer.articles
        if article.get('date') == date
    ]

    if not articles:
        if progress_callback:
            progress_callback(0, "لم يتم العثور على مقالات", "error")
        return None, None

    total_steps = len(articles) * 2 + 2  # per article: content fetch + formatting (2 steps) + init + save
    current_step = 0

    def update_progress(step_increment, message, status="processing"):
        nonlocal current_step
        current_step += step_increment
        percentage = int((current_step / total_steps) * 100) if total_steps > 0 else 0
        if progress_callback:
            progress_callback(percentage, message, status)

    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Inches(0.12)

    update_progress(1, f"بدء إنشاء المستند ({len(articles)} مقال)", "processing")
    
    title = doc.add_heading(f'أخبار فلسطين - {analyzer.format_date_arabic(date)}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for idx, article in enumerate(articles, start=1):
        update_progress(0, f"معالجة المقال {idx} من {len(articles)}: {article.get('title', 'بدون عنوان')[:50]}...", "processing")
        
        heading = doc.add_heading(article.get('title', 'بدون عنوان'), level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add link (same format as with images - decoded URL, no label)
        if article.get('link'):
            cleaned_url = unquote(article['link'])
            link_para = doc.add_paragraph()
            link_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = link_para.add_run(cleaned_url)
            run.font.size = Inches(0.11)

        # Fetch and add content (using the same formatted function as with images)
        content_text = ''
        update_progress(0, f"جاري جلب محتوى المقال {idx}...", "processing")
        content_text = fetch_article_content_formatted(article)
        if not content_text and article.get('excerpt'):
            content_text = article.get('excerpt')

        if not content_text:
            content_text = 'عذراً، تعذر تحميل المحتوى الكامل للمقال.'

        for paragraph in content_text.split('\n'):
            cleaned = paragraph.strip()
            if not cleaned:
                continue
            para = doc.add_paragraph(cleaned)
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            para.paragraph_format.keep_together = True

        update_progress(1, f"تمت إضافة محتوى المقال {idx}", "processing")

        if idx != len(articles):
            doc.add_paragraph('-' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
 
    update_progress(1, "جاري حفظ الملف...", "processing")
    footer = doc.add_paragraph(f'تم إنشاء هذا التقرير في: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    filename = f"palestine_news_full_content_{date.replace('-', '_')}.docx"
    filepath = TEMP_DIR / filename
    doc.save(filepath)

    update_progress(1, "اكتمل إنشاء الملف بنجاح!", "completed")
    return filepath, filename


def create_document_with_images(date: str, include_content: bool = True) -> Tuple[Optional[Path], Optional[str]]:
    articles = [
        article for article in analyzer.articles
        if article.get('date') == date
    ]

    if not articles:
        return None, None

    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Inches(0.12)

    title = doc.add_heading(f'أخبار فلسطين - {analyzer.format_date_arabic(date)}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Track image files to clean up after document creation
    used_image_paths = []

    for idx, article in enumerate(articles, start=1):
        heading = doc.add_heading(article.get('title', 'بدون عنوان'), level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        image_path = ensure_image_file(article)
        if image_path and image_path.exists():
            try:
                resized_stream = prepare_image_stream(image_path)
                if resized_stream:
                    doc.add_picture(resized_stream, width=Inches(5.5))
                else:
                    doc.add_picture(str(image_path), width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Track this image for cleanup
                used_image_paths.append(image_path)
            except Exception as exc:
                print(f"Failed to insert image {image_path}: {exc}")

        if article.get('link'):
            cleaned_url = unquote(article['link'])
            link_para = doc.add_paragraph()
            link_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = link_para.add_run(cleaned_url)
            run.font.size = Inches(0.11)

        content_text = ''
        if include_content:
            content_text = fetch_article_content_formatted(article)
            if not content_text and article.get('excerpt'):
                content_text = article.get('excerpt')
        else:
            content_text = article.get('excerpt') or ''

        if not content_text:
            content_text = 'عذراً، تعذر تحميل المحتوى الكامل للمقال.'

        for paragraph in content_text.split('\n'):
            cleaned = paragraph.strip()
            if not cleaned:
                continue
            para = doc.add_paragraph(cleaned)
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            para.paragraph_format.keep_together = True

        if idx != len(articles):
            doc.add_paragraph('-' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
 
    footer = doc.add_paragraph(f'تم إنشاء هذا التقرير في: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
 
    # Use simpler filename format (matching other exports)
    filename = f"palestine_news_with_images_{date.replace('-', '_')}.docx"
    filepath = TEMP_DIR / filename
    doc.save(filepath)

    # Clean up cached images after embedding (Option 2)
    for image_path in used_image_paths:
        try:
            if image_path.exists():
                image_path.unlink()
                print(f"Cleaned up cached image: {image_path}")
        except Exception as exc:
            print(f"Failed to delete cached image {image_path}: {exc}")

    return filepath, filename


def create_document_with_images_progress(date: str, include_content: bool = True, progress_callback: Optional[Callable] = None) -> Tuple[Optional[Path], Optional[str]]:
    """Create document with images, reporting progress via callback"""
    articles = [
        article for article in analyzer.articles
        if article.get('date') == date
    ]

    if not articles:
        if progress_callback:
            progress_callback(0, "لم يتم العثور على مقالات", "error")
        return None, None

    total_steps = len(articles) * 3 + 3  # per article: image, content, formatting (3 steps) + init (1) + save (1) + cleanup (1)
    current_step = 0

    def update_progress(step_increment, message, status="processing"):
        nonlocal current_step
        current_step += step_increment
        percentage = int((current_step / total_steps) * 100)
        if progress_callback:
            progress_callback(percentage, message, status)

    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Inches(0.12)

    update_progress(1, f"بدء إنشاء المستند ({len(articles)} مقال)", "processing")
    
    title = doc.add_heading(f'أخبار فلسطين - {analyzer.format_date_arabic(date)}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Track image files to clean up after document creation
    used_image_paths = []

    for idx, article in enumerate(articles, start=1):
        update_progress(0, f"معالجة المقال {idx} من {len(articles)}: {article.get('title', 'بدون عنوان')[:50]}...", "processing")
        
        heading = doc.add_heading(article.get('title', 'بدون عنوان'), level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Download and process image
        image_path = ensure_image_file(article)
        if image_path and image_path.exists():
            try:
                update_progress(0, f"إضافة صورة للمقال {idx}...", "processing")
                resized_stream = prepare_image_stream(image_path)
                if resized_stream:
                    doc.add_picture(resized_stream, width=Inches(5.5))
                else:
                    doc.add_picture(str(image_path), width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                used_image_paths.append(image_path)
            except Exception as exc:
                print(f"Failed to insert image {image_path}: {exc}")
        update_progress(1, f"تمت إضافة صورة المقال {idx}", "processing")

        # Add link
        if article.get('link'):
            cleaned_url = unquote(article['link'])
            link_para = doc.add_paragraph()
            link_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = link_para.add_run(cleaned_url)
            run.font.size = Inches(0.11)

        # Fetch and add content
        content_text = ''
        if include_content:
            update_progress(0, f"جاري جلب محتوى المقال {idx}...", "processing")
            content_text = fetch_article_content_formatted(article)
            if not content_text and article.get('excerpt'):
                content_text = article.get('excerpt')
        else:
            content_text = article.get('excerpt') or ''

        if not content_text:
            content_text = 'عذراً، تعذر تحميل المحتوى الكامل للمقال.'

        update_progress(1, f"إضافة محتوى المقال {idx}...", "processing")
        for paragraph in content_text.split('\n'):
            cleaned = paragraph.strip()
            if not cleaned:
                continue
            para = doc.add_paragraph(cleaned)
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            para.paragraph_format.keep_together = True

        if idx != len(articles):
            doc.add_paragraph('-' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
        
        update_progress(1, f"اكتمل المقال {idx}", "processing")
 
    update_progress(1, "جاري حفظ الملف...", "processing")
    footer = doc.add_paragraph(f'تم إنشاء هذا التقرير في: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
 
    # Use simpler filename format (matching other exports)
    filename = f"palestine_news_with_images_{date.replace('-', '_')}.docx"
    filepath = TEMP_DIR / filename
    doc.save(filepath)

    # Clean up cached images after embedding (Option 2)
    update_progress(0, "جاري تنظيف الصور المؤقتة...", "processing")
    for image_path in used_image_paths:
        try:
            if image_path.exists():
                image_path.unlink()
        except Exception as exc:
            print(f"Failed to delete cached image {image_path}: {exc}")
    
    update_progress(1, "اكتمل إنشاء الملف بنجاح!", "completed")
    return filepath, filename


@app.route('/api/export/headline-only')
def api_export_headline_only():
    date = request.args.get('date', '').strip()

    if not date:
        return jsonify({'error': 'Date parameter is required'}), 400

    filepath, filename = create_headline_only_document(date)

    if not filepath:
        return jsonify({'error': 'No articles found for the specified date'}), 404

    return send_file(filepath, as_attachment=True, download_name=filename)


@app.route('/api/export/word-with-images')
def api_export_with_images():
    date = request.args.get('date', '').strip()
    include_content = request.args.get('include_content', 'true').lower() != 'false'

    if not date:
        return jsonify({'error': 'Date parameter is required'}), 400

    filepath, filename = create_document_with_images(date, include_content)

    if not filepath:
        return jsonify({'error': 'No articles found for the specified date'}), 404

    # Convert Path to absolute string path for Flask send_file
    if isinstance(filepath, Path):
        filepath_str = str(filepath.resolve())
    else:
        filepath_str = os.path.abspath(filepath)
    
    # Verify file exists
    if not os.path.exists(filepath_str):
        return jsonify({'error': 'Generated file not found'}), 500
    
    # Ensure filename is properly encoded for Content-Disposition header
    try:
        response = send_file(
            filepath_str, 
            as_attachment=True, 
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
        # Explicitly set Content-Disposition header to ensure correct filename
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"; filename*=UTF-8\'\'{quote(filename)}'
        
        return response
    except Exception as e:
        print(f"Error sending file: {e}")
        return jsonify({'error': f'Failed to send file: {str(e)}'}), 500


@app.route('/api/export/word-with-images/start', methods=['POST'])
def api_export_with_images_start():
    """Start document creation job and return job_id"""
    data = request.get_json()
    date = data.get('date', '').strip()
    include_content = data.get('include_content', True)

    if not date:
        return jsonify({'error': 'Date parameter is required'}), 400

    # Generate unique job ID
    job_id = str(uuid.uuid4())
    
    # Initialize progress
    with progress_lock:
        progress_tracker[job_id] = {
            'percentage': 0,
            'message': 'بدء العملية...',
            'status': 'processing',
            'filepath': None,
            'filename': None,
            'error': None
        }

    # Start document creation in background thread
    def create_document_thread():
        def progress_callback(percentage, message, status):
            with progress_lock:
                if job_id in progress_tracker:
                    progress_tracker[job_id]['percentage'] = percentage
                    progress_tracker[job_id]['message'] = message
                    progress_tracker[job_id]['status'] = status

        try:
            filepath, filename = create_document_with_images_progress(
                date, include_content, progress_callback
            )
            with progress_lock:
                if job_id in progress_tracker:
                    if filepath:
                        progress_tracker[job_id]['filepath'] = str(filepath.resolve()) if isinstance(filepath, Path) else filepath
                        progress_tracker[job_id]['filename'] = filename
                        progress_tracker[job_id]['status'] = 'completed'
                        progress_tracker[job_id]['percentage'] = 100
                        progress_tracker[job_id]['message'] = 'اكتمل إنشاء الملف بنجاح!'
                    else:
                        progress_tracker[job_id]['status'] = 'error'
                        progress_tracker[job_id]['error'] = 'لم يتم العثور على مقالات'
        except Exception as e:
            with progress_lock:
                if job_id in progress_tracker:
                    progress_tracker[job_id]['status'] = 'error'
                    progress_tracker[job_id]['error'] = str(e)
                    progress_tracker[job_id]['message'] = f'حدث خطأ: {str(e)}'

    thread = threading.Thread(target=create_document_thread)
    thread.daemon = True
    thread.start()

    return jsonify({'job_id': job_id})


@app.route('/api/export/word-with-images/progress/<job_id>')
def api_export_with_images_progress(job_id):
    """SSE endpoint for streaming progress updates"""
    def generate():
        last_percentage = -1
        while True:
            with progress_lock:
                progress = progress_tracker.get(job_id, {})
                percentage = progress.get('percentage', 0)
                message = progress.get('message', '')
                status = progress.get('status', 'processing')
                filepath = progress.get('filepath')
                filename = progress.get('filename')
                error = progress.get('error')

            # Only send update if percentage changed or status changed
            if percentage != last_percentage or status in ['completed', 'error']:
                data = {
                    'percentage': percentage,
                    'message': message,
                    'status': status
                }
                
                if status == 'completed' and filepath and filename:
                    data['filepath'] = filepath
                    data['filename'] = filename
                    data['download_url'] = f'/api/export/word-with-images/download/{job_id}'
                
                if status == 'error' and error:
                    data['error'] = error

                yield f"data: {json.dumps(data, ensure_ascii=False)}\n\n"
                last_percentage = percentage

                if status in ['completed', 'error']:
                    # Clean up after a delay
                    time.sleep(2)
                    with progress_lock:
                        if job_id in progress_tracker:
                            del progress_tracker[job_id]
                    break

            time.sleep(0.5)  # Update every 500ms

    return Response(
        stream_with_context(generate()),
        mimetype='text/event-stream',
        headers={
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no',
            'Connection': 'keep-alive'
        }
    )


@app.route('/api/export/word-with-images/download/<job_id>')
def api_export_with_images_download(job_id):
    """Download the completed document"""
    with progress_lock:
        progress = progress_tracker.get(job_id, {})
        filepath = progress.get('filepath')
        filename = progress.get('filename')

    if not filepath or not filename:
        return jsonify({'error': 'File not found or job not completed'}), 404

    if not os.path.exists(filepath):
        return jsonify({'error': 'Generated file not found'}), 404

    try:
        response = send_file(
            filepath,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"; filename*=UTF-8\'\'{quote(filename)}'
        return response
    except Exception as e:
        print(f"Error sending file: {e}")
        return jsonify({'error': f'Failed to send file: {str(e)}'}), 500


@app.route('/api/export/word/start', methods=['POST'])
def api_export_word_start():
    """Start document creation job (with summaries) and return job_id"""
    data = request.get_json()
    date = data.get('date', '').strip()
    include_content = data.get('include_content', False)

    if not date:
        return jsonify({'error': 'Date parameter is required'}), 400

    articles = [a for a in analyzer.articles if a.get('date') == date]
    if not articles:
        return jsonify({'error': 'No articles found for the specified date'}), 404

    job_id = str(uuid.uuid4())
    
    with progress_lock:
        progress_tracker[job_id] = {
            'percentage': 0,
            'message': 'بدء العملية...',
            'status': 'processing',
            'filepath': None,
            'filename': None,
            'error': None
        }

    def create_document_thread():
        def progress_callback(percentage, message, status):
            with progress_lock:
                if job_id in progress_tracker:
                    progress_tracker[job_id]['percentage'] = percentage
                    progress_tracker[job_id]['message'] = message
                    progress_tracker[job_id]['status'] = status

        try:
            doc = analyzer.create_word_document(articles, date, include_content, progress_callback)
            filename = f"palestine_news_with_summaries_{date.replace('-', '_')}.docx"
            filepath = TEMP_DIR / filename
            doc.save(filepath)
            
            with progress_lock:
                if job_id in progress_tracker:
                    progress_tracker[job_id]['filepath'] = str(filepath.resolve())
                    progress_tracker[job_id]['filename'] = filename
                    progress_tracker[job_id]['status'] = 'completed'
                    progress_tracker[job_id]['percentage'] = 100
                    progress_tracker[job_id]['message'] = 'اكتمل إنشاء الملف بنجاح!'
        except Exception as e:
            with progress_lock:
                if job_id in progress_tracker:
                    progress_tracker[job_id]['status'] = 'error'
                    progress_tracker[job_id]['error'] = str(e)
                    progress_tracker[job_id]['message'] = f'حدث خطأ: {str(e)}'

    thread = threading.Thread(target=create_document_thread)
    thread.daemon = True
    thread.start()

    return jsonify({'job_id': job_id})


@app.route('/api/export/headline-only/start', methods=['POST'])
def api_export_headline_only_start():
    """Start headline-only document creation job and return job_id"""
    data = request.get_json()
    date = data.get('date', '').strip()

    if not date:
        return jsonify({'error': 'Date parameter is required'}), 400

    job_id = str(uuid.uuid4())
    
    with progress_lock:
        progress_tracker[job_id] = {
            'percentage': 0,
            'message': 'بدء العملية...',
            'status': 'processing',
            'filepath': None,
            'filename': None,
            'error': None
        }

    def create_document_thread():
        def progress_callback(percentage, message, status):
            with progress_lock:
                if job_id in progress_tracker:
                    progress_tracker[job_id]['percentage'] = percentage
                    progress_tracker[job_id]['message'] = message
                    progress_tracker[job_id]['status'] = status

        try:
            filepath, filename = create_headline_only_document(date, progress_callback)
            with progress_lock:
                if job_id in progress_tracker:
                    if filepath:
                        progress_tracker[job_id]['filepath'] = str(filepath.resolve()) if isinstance(filepath, Path) else filepath
                        progress_tracker[job_id]['filename'] = filename
                        progress_tracker[job_id]['status'] = 'completed'
                        progress_tracker[job_id]['percentage'] = 100
                        progress_tracker[job_id]['message'] = 'اكتمل إنشاء الملف بنجاح!'
                    else:
                        progress_tracker[job_id]['status'] = 'error'
                        progress_tracker[job_id]['error'] = 'لم يتم العثور على مقالات'
        except Exception as e:
            with progress_lock:
                if job_id in progress_tracker:
                    progress_tracker[job_id]['status'] = 'error'
                    progress_tracker[job_id]['error'] = str(e)
                    progress_tracker[job_id]['message'] = f'حدث خطأ: {str(e)}'

    thread = threading.Thread(target=create_document_thread)
    thread.daemon = True
    thread.start()

    return jsonify({'job_id': job_id})


@app.route('/api/export/word/progress/<job_id>')
def api_export_word_progress(job_id):
    """SSE endpoint for streaming progress updates (word with summaries)"""
    def generate():
        last_percentage = -1
        while True:
            with progress_lock:
                progress = progress_tracker.get(job_id, {})
                percentage = progress.get('percentage', 0)
                message = progress.get('message', '')
                status = progress.get('status', 'processing')
                filepath = progress.get('filepath')
                filename = progress.get('filename')
                error = progress.get('error')

            if percentage != last_percentage or status in ['completed', 'error']:
                data = {
                    'percentage': percentage,
                    'message': message,
                    'status': status
                }
                
                if status == 'completed' and filepath and filename:
                    data['filepath'] = filepath
                    data['filename'] = filename
                    data['download_url'] = f'/api/export/word/download/{job_id}'
                
                if status == 'error' and error:
                    data['error'] = error

                yield f"data: {json.dumps(data, ensure_ascii=False)}\n\n"
                last_percentage = percentage

                if status in ['completed', 'error']:
                    time.sleep(2)
                    with progress_lock:
                        if job_id in progress_tracker:
                            del progress_tracker[job_id]
                    break

            time.sleep(0.5)

    return Response(
        stream_with_context(generate()),
        mimetype='text/event-stream',
        headers={
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no',
            'Connection': 'keep-alive'
        }
    )


@app.route('/api/export/headline-only/progress/<job_id>')
def api_export_headline_only_progress(job_id):
    """SSE endpoint for streaming progress updates (headline only)"""
    def generate():
        last_percentage = -1
        while True:
            with progress_lock:
                progress = progress_tracker.get(job_id, {})
                percentage = progress.get('percentage', 0)
                message = progress.get('message', '')
                status = progress.get('status', 'processing')
                filepath = progress.get('filepath')
                filename = progress.get('filename')
                error = progress.get('error')

            if percentage != last_percentage or status in ['completed', 'error']:
                data = {
                    'percentage': percentage,
                    'message': message,
                    'status': status
                }
                
                if status == 'completed' and filepath and filename:
                    data['filepath'] = filepath
                    data['filename'] = filename
                    data['download_url'] = f'/api/export/headline-only/download/{job_id}'
                
                if status == 'error' and error:
                    data['error'] = error

                yield f"data: {json.dumps(data, ensure_ascii=False)}\n\n"
                last_percentage = percentage

                if status in ['completed', 'error']:
                    time.sleep(2)
                    with progress_lock:
                        if job_id in progress_tracker:
                            del progress_tracker[job_id]
                    break

            time.sleep(0.5)

    return Response(
        stream_with_context(generate()),
        mimetype='text/event-stream',
        headers={
            'Cache-Control': 'no-cache',
            'X-Accel-Buffering': 'no',
            'Connection': 'keep-alive'
        }
    )


@app.route('/api/export/word/download/<job_id>')
def api_export_word_download(job_id):
    """Download the completed document (word with summaries)"""
    with progress_lock:
        progress = progress_tracker.get(job_id, {})
        filepath = progress.get('filepath')
        filename = progress.get('filename')

    if not filepath or not filename:
        return jsonify({'error': 'File not found or job not completed'}), 404

    if not os.path.exists(filepath):
        return jsonify({'error': 'Generated file not found'}), 404

    try:
        response = send_file(
            filepath,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"; filename*=UTF-8\'\'{quote(filename)}'
        return response
    except Exception as e:
        print(f"Error sending file: {e}")
        return jsonify({'error': f'Failed to send file: {str(e)}'}), 500


@app.route('/api/export/headline-only/download/<job_id>')
def api_export_headline_only_download(job_id):
    """Download the completed document (headline only)"""
    with progress_lock:
        progress = progress_tracker.get(job_id, {})
        filepath = progress.get('filepath')
        filename = progress.get('filename')

    if not filepath or not filename:
        return jsonify({'error': 'File not found or job not completed'}), 404

    if not os.path.exists(filepath):
        return jsonify({'error': 'Generated file not found'}), 404

    try:
        response = send_file(
            filepath,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"; filename*=UTF-8\'\'{quote(filename)}'
        return response
    except Exception as e:
        print(f"Error sending file: {e}")
        return jsonify({'error': f'Failed to send file: {str(e)}'}), 500


@app.route('/api/article/<int:article_id>/content')
def api_article_content(article_id):
    """API endpoint to fetch full article content"""
    article = next((a for a in analyzer.articles if a.get('id') == article_id), None)
    if not article:
        return jsonify({'error': 'Article not found'}), 404
    
    article_url = article.get('link')
    if not article_url:
        return jsonify({'error': 'Article URL not found'}), 404
    
    content = analyzer.fetch_article_content(article_url)
    return jsonify({'content': content})

@app.route('/api/export/word')
def api_export_word():
    """API endpoint to export articles to Word document"""
    date = request.args.get('date', '')
    include_content = request.args.get('include_content', 'false').lower() == 'true'
    
    if not date:
        return jsonify({'error': 'Date parameter is required'}), 400
    
    # Get articles for the specified date
    articles = []
    for article in analyzer.articles:
        if article.get('date') == date:
            articles.append(article)
    
    if not articles:
        return jsonify({'error': 'No articles found for the specified date'}), 404
    
    try:
        # Create Word document
        doc = analyzer.create_word_document(articles, date, include_content)
        
        # Save to temporary file
        filename = f"palestine_news_with_summaries_{date.replace('-', '_')}.docx"
        temp_path = os.path.join('temp', filename)
        
        # Create temp directory if it doesn't exist
        os.makedirs('temp', exist_ok=True)
        
        doc.save(temp_path)
        
        return send_file(temp_path, as_attachment=True, download_name=filename)
        
    except Exception as e:
        print(f"Error creating Word document: {e}")
        return jsonify({'error': 'Failed to create Word document'}), 500

@app.route('/article/<int:article_id>')
def article_detail(article_id):
    """Article detail page"""
    lang = get_language()
    article = next((a for a in analyzer.articles if a.get('id') == article_id), None)
    if not article:
        return "Article not found", 404
    return render_template('article_detail.html', article=article, lang=lang)

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
