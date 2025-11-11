#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Create a Word document with article headlines followed by full content
(no excerpt/summary section).
"""

import json
import re
import time
from datetime import datetime
from urllib.parse import unquote

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


def format_date_arabic(date_str: str) -> str:
    """Return a date string formatted using Arabic month names."""
    try:
        date_obj = datetime.strptime(date_str, "%Y-%m-%d")
        arabic_months = [
            "يناير",
            "فبراير",
            "مارس",
            "أبريل",
            "مايو",
            "يونيو",
            "يوليو",
            "أغسطس",
            "سبتمبر",
            "أكتوبر",
            "نوفمبر",
            "ديسمبر",
        ]
        return f"{date_obj.day} {arabic_months[date_obj.month - 1]} {date_obj.year}"
    except Exception:
        return date_str


def decode_url(url: str) -> str:
    """Decode URL-encoded Arabic characters for display."""
    try:
        return unquote(url, encoding="utf-8")
    except Exception:
        return url


def fetch_article_content(article_url: str) -> str:
    """Fetch and clean the full text content of an article."""
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
            "(KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
            "Accept-Language": "ar,en-US;q=0.7,en;q=0.3",
            "Connection": "keep-alive",
        }
        response = requests.get(article_url, headers=headers, timeout=15)
        response.raise_for_status()

        soup = BeautifulSoup(response.content, "html.parser")

        selectors = [
            ".wysiwyg--all-content",
            ".article-body",
            ".post-content",
            "article .content",
            ".entry-content",
        ]
        article_content = None
        for selector in selectors:
            node = soup.select_one(selector)
            if node:
                article_content = node
                break

        if not article_content:
            fallback = soup.find_all(
                "div", string=lambda text: text and len(text.strip()) > 200
            )
            if fallback:
                article_content = fallback[0]

        if not article_content:
            return "عذراً، لم يتم العثور على المحتوى الكامل للمقال." 

        # Remove unwanted elements and inline link tags while preserving text
        for unwanted in article_content.find_all([
            "script",
            "style",
            "nav",
            "header",
            "footer",
            "aside",
        ]):
            unwanted.decompose()

        for link in article_content.find_all("a"):
            link.replace_with(link.get_text())

        # Extract text with natural paragraph boundaries
        text = article_content.get_text(separator=" \n ", strip=True)

        # Normalize whitespace: collapse spaces but keep sentence flow
        text = re.sub(r"\s+", " ", text)

        # Restore paragraph breaks using punctuation hints (Arabic/Latin)
        paragraph_candidates = re.split(r"([.!؟]\s+)", text)
        paragraphs, buffer = [], []
        for part in paragraph_candidates:
            if part.strip():
                buffer.append(part)
                if part.strip().endswith(tuple(".!؟")) and len(buffer) >= 2:
                    paragraphs.append("".join(buffer).strip())
                    buffer = []
        if buffer:
            paragraphs.append("".join(buffer).strip())

        # Fallback if splitting failed
        if not paragraphs:
            paragraphs = [text]

        return "\n\n".join(paragraphs)

    except requests.RequestException as exc:
        print(f"⚠️  Network error while fetching article: {exc}")
        return "عذراً، حدث خطأ أثناء تحميل محتوى المقال." 
    except Exception as exc:
        print(f"⚠️  Parsing error: {exc}")
        return "عذراً، حدث خطأ أثناء معالجة محتوى المقال." 


def create_document(articles: list, date: str) -> Document:
    """Create a Word document with each article headline followed by full content."""
    doc = Document()
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Inches(0.12)

    title = doc.add_heading(f"أخبار فلسطين - {format_date_arabic(date)}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for idx, article in enumerate(articles, start=1):
        print(f"Processing article {idx}/{len(articles)}: {article.get('title', '')[:60]}...")

        heading = doc.add_heading(article.get("title", "بدون عنوان"), level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if article.get("link"):
            cleaned_url = decode_url(article["link"])
            link_para = doc.add_paragraph()
            link_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = link_para.add_run(cleaned_url)
            run.font.size = Inches(0.11)

        if article.get("link"):
            content_text = fetch_article_content(article["link"])
        else:
            content_text = "عذراً، رابط المقال غير متوفر." 

        for paragraph in content_text.split("\n"):
            cleaned = paragraph.strip()
            if not cleaned:
                continue
            p = doc.add_paragraph(cleaned)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.keep_together = True

        if idx != len(articles):
            doc.add_paragraph("-" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

        time.sleep(0.4)  # gentle delay between requests

    footer = doc.add_paragraph(
        f"تم إنشاء هذا التقرير في: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


def main():
    date = "2025-11-06"  # TODO: Change to desired date
    print(f"Loading articles for date {date}...")

    with open("articles_combined.json", "r", encoding="utf-8") as fh:
        all_articles = json.load(fh)

    articles = [art for art in all_articles if art.get("date") == date]
    if not articles:
        print("❌ No articles found for this date.")
        return

    document = create_document(articles, date)
    timestamp = datetime.now().strftime("%H%M%S")
    filename = f"palestine_news_full_headline_only_{date.replace('-', '_')}_{timestamp}.docx"
    document.save(filename)

    print("✅ Document created successfully!")
    print(f"📄 Articles included: {len(articles)}")
    print(f"📁 Saved as: {filename}")


if __name__ == "__main__":
    main()
